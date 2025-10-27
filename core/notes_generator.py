"""
StudyMate Notes Generator - AI-Powered Note Creation & Export
Generates structured notes, summaries, and exports to multiple formats
"""

import os
import uuid
from typing import Dict, List, Optional, Any
from datetime import datetime
from pathlib import Path
import json
import markdown
import html2text

# Optional imports for export
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class NotesGenerator:
    """Advanced notes generation with multiple export formats"""
    
    def __init__(self):
        print("📝 Initializing Notes Generator...")
        
        # Configuration
        self.export_dir = Path("exports")
        self.export_dir.mkdir(exist_ok=True)
        
        # Export capabilities
        self.pdf_available = PDF_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        # Note templates
        self.templates = {
            "basic": {
                "structure": ["Introduction", "Main Content", "Key Points", "Summary"],
                "style": "simple"
            },
            "academic": {
                "structure": ["Overview", "Detailed Analysis", "Examples", "Applications", "Conclusion"],
                "style": "formal"
            },
            "study_guide": {
                "structure": ["Learning Objectives", "Key Concepts", "Important Facts", "Practice Questions", "Review"],
                "style": "educational"
            },
            "summary": {
                "structure": ["Executive Summary", "Main Points", "Details", "Takeaways"],
                "style": "concise"
            }
        }
        
        print(f"✅ Notes Generator ready - PDF: {self.pdf_available}, DOCX: {self.docx_available}")
    
    async def generate_notes(
        self,
        topic: str,
        user_id: str,
        context: Optional[Dict] = None,
        gemini_model = None
    ) -> str:
        """Generate basic notes for a topic"""
        
        if gemini_model:
            return await self._generate_ai_notes(topic, gemini_model, "basic")
        else:
            return await self._generate_template_notes(topic, "basic")
    
    async def generate_comprehensive_notes(
        self,
        topic: str,
        user_id: str,
        format_type: str = "markdown",
        detail_level: str = "medium",
        include_examples: bool = True,
        template: str = "academic",
        gemini_model = None
    ) -> str:
        """Generate comprehensive structured notes"""
        
        try:
            if gemini_model:
                notes_content = await self._generate_ai_notes(
                    topic, gemini_model, template, detail_level, include_examples
                )
            else:
                notes_content = await self._generate_template_notes(topic, template)
            
            # Format according to requested type
            if format_type == "html":
                return markdown.markdown(notes_content)
            elif format_type == "plain":
                return html2text.html2text(markdown.markdown(notes_content))
            else:  # markdown (default)
                return notes_content
                
        except Exception as e:
            print(f"❌ Error generating comprehensive notes: {e}")
            return f"# {topic}\n\nError generating notes: {str(e)}"
    
    async def _generate_ai_notes(
        self,
        topic: str,
        gemini_model,
        template: str = "academic",
        detail_level: str = "medium",
        include_examples: bool = True
    ) -> str:
        """Generate AI-powered notes using Gemini"""
        
        try:
            template_info = self.templates.get(template, self.templates["academic"])
            structure = template_info["structure"]
            
            # Create comprehensive prompt
            prompt = f"""Generate comprehensive study notes on: {topic}

Structure the notes with these sections:
{chr(10).join([f"- {section}" for section in structure])}

Requirements:
- Detail level: {detail_level}
- Include examples: {include_examples}
- Format: Well-structured markdown with headers, bullet points, and emphasis
- Style: {template_info['style']}
- Length: {"Detailed" if detail_level == "high" else "Moderate" if detail_level == "medium" else "Concise"}

Make the notes educational, clear, and well-organized for effective studying."""

            if include_examples:
                prompt += "\n- Include relevant examples and practical applications"
            
            if detail_level == "high":
                prompt += "\n- Provide in-depth explanations and comprehensive coverage"
            elif detail_level == "low":
                prompt += "\n- Keep explanations concise and focus on key points"
            
            # Generate notes
            response = gemini_model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            print(f"❌ Error generating AI notes: {e}")
            return await self._generate_template_notes(topic, template)
    
    async def _generate_template_notes(self, topic: str, template: str = "basic") -> str:
        """Generate template-based notes when AI is not available"""
        
        template_info = self.templates.get(template, self.templates["basic"])
        structure = template_info["structure"]
        
        notes = f"# {topic}\n\n"
        notes += f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n"
        
        for section in structure:
            notes += f"## {section}\n\n"
            
            if section.lower() in ["introduction", "overview"]:
                notes += f"This section provides an introduction to {topic}.\n\n"
            elif section.lower() in ["main content", "detailed analysis"]:
                notes += f"Key concepts and detailed information about {topic}:\n\n"
                notes += "- Important point 1\n"
                notes += "- Important point 2\n"
                notes += "- Important point 3\n\n"
            elif section.lower() in ["key points", "key concepts"]:
                notes += f"Essential points to remember about {topic}:\n\n"
                notes += "1. First key concept\n"
                notes += "2. Second key concept\n"
                notes += "3. Third key concept\n\n"
            elif section.lower() in ["examples", "applications"]:
                notes += f"Practical examples and applications of {topic}:\n\n"
                notes += "- Example 1: [Description]\n"
                notes += "- Example 2: [Description]\n\n"
            elif section.lower() in ["summary", "conclusion", "takeaways"]:
                notes += f"Summary of key learnings about {topic}:\n\n"
                notes += f"In conclusion, {topic} is an important concept that...\n\n"
            else:
                notes += f"Content for {section.lower()} goes here.\n\n"
        
        notes += "---\n\n"
        notes += "*Note: For AI-powered detailed notes, please configure your Gemini API key.*\n"
        
        return notes
    
    async def export_notes(
        self,
        content: str,
        format_type: str,
        filename: str,
        metadata: Optional[Dict] = None
    ) -> str:
        """Export notes to specified format"""
        
        try:
            # Ensure filename has correct extension
            base_name = Path(filename).stem
            
            if format_type.lower() == "pdf":
                return await self._export_to_pdf(content, f"{base_name}.pdf", metadata)
            elif format_type.lower() == "docx":
                return await self._export_to_docx(content, f"{base_name}.docx", metadata)
            elif format_type.lower() == "html":
                return await self._export_to_html(content, f"{base_name}.html", metadata)
            elif format_type.lower() == "txt":
                return await self._export_to_txt(content, f"{base_name}.txt", metadata)
            else:  # markdown (default)
                return await self._export_to_markdown(content, f"{base_name}.md", metadata)
                
        except Exception as e:
            print(f"❌ Error exporting notes: {e}")
            raise
    
    async def _export_to_pdf(self, content: str, filename: str, metadata: Optional[Dict] = None) -> str:
        """Export notes to PDF format"""
        
        if not self.pdf_available:
            raise Exception("PDF export not available. Install reportlab: pip install reportlab")
        
        try:
            file_path = self.export_dir / filename
            
            # Create PDF document
            doc = SimpleDocTemplate(str(file_path), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            # Parse markdown content
            lines = content.split('\n')
            current_text = ""
            
            for line in lines:
                line = line.strip()
                
                if line.startswith('# '):
                    # Main title
                    if current_text:
                        story.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    story.append(Paragraph(line[2:], title_style))
                    story.append(Spacer(1, 12))
                    
                elif line.startswith('## '):
                    # Section header
                    if current_text:
                        story.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    story.append(Paragraph(line[3:], styles['Heading2']))
                    story.append(Spacer(1, 6))
                    
                elif line.startswith('### '):
                    # Subsection header
                    if current_text:
                        story.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    story.append(Paragraph(line[4:], styles['Heading3']))
                    story.append(Spacer(1, 6))
                    
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    if current_text:
                        story.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                    
                elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                    # Numbered list
                    if current_text:
                        story.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    story.append(Paragraph(line, styles['Normal']))
                    
                elif line == "---":
                    # Horizontal rule
                    if current_text:
                        story.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    story.append(Spacer(1, 12))
                    
                elif line:
                    # Regular text
                    current_text += line + " "
                    
                else:
                    # Empty line - paragraph break
                    if current_text:
                        story.append(Paragraph(current_text.strip(), styles['Normal']))
                        story.append(Spacer(1, 6))
                        current_text = ""
            
            # Add remaining text
            if current_text:
                story.append(Paragraph(current_text.strip(), styles['Normal']))
            
            # Add metadata footer
            if metadata:
                story.append(Spacer(1, 20))
                footer_text = f"Generated by StudyMate AI on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                story.append(Paragraph(footer_text, styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            print(f"📄 Exported PDF: {filename}")
            return str(file_path)
            
        except Exception as e:
            print(f"❌ PDF export error: {e}")
            raise
    
    async def _export_to_docx(self, content: str, filename: str, metadata: Optional[Dict] = None) -> str:
        """Export notes to DOCX format"""
        
        if not self.docx_available:
            raise Exception("DOCX export not available. Install python-docx: pip install python-docx")
        
        try:
            file_path = self.export_dir / filename
            
            # Create DOCX document
            doc = DocxDocument()
            
            # Parse markdown content
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if line.startswith('# '):
                    # Main title
                    heading = doc.add_heading(line[2:], level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                elif line.startswith('## '):
                    # Section header
                    doc.add_heading(line[3:], level=2)
                    
                elif line.startswith('### '):
                    # Subsection header
                    doc.add_heading(line[4:], level=3)
                    
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.add_run(line[2:])
                    
                elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                    # Numbered list
                    p = doc.add_paragraph()
                    p.style = 'List Number'
                    p.add_run(line[3:])
                    
                elif line == "---":
                    # Horizontal rule (add space)
                    doc.add_paragraph()
                    
                elif line.startswith('*') and line.endswith('*'):
                    # Italic text
                    p = doc.add_paragraph()
                    run = p.add_run(line[1:-1])
                    run.italic = True
                    
                elif line:
                    # Regular paragraph
                    doc.add_paragraph(line)
            
            # Add metadata
            if metadata:
                doc.add_paragraph()
                footer = doc.add_paragraph(f"Generated by StudyMate AI on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(str(file_path))
            
            print(f"📄 Exported DOCX: {filename}")
            return str(file_path)
            
        except Exception as e:
            print(f"❌ DOCX export error: {e}")
            raise
    
    async def _export_to_html(self, content: str, filename: str, metadata: Optional[Dict] = None) -> str:
        """Export notes to HTML format"""
        
        try:
            file_path = self.export_dir / filename
            
            # Convert markdown to HTML
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # Create full HTML document
            html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>StudyMate Notes</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        h1 {{ border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }}
        code {{ background: #f8f9fa; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background: #f8f9fa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #3498db; margin: 0; padding-left: 20px; color: #7f8c8d; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #eee; text-align: center; color: #7f8c8d; font-size: 0.9em; }}
    </style>
</head>
<body>
    {html_content}
    {f'<div class="footer">Generated by StudyMate AI on {datetime.now().strftime("%Y-%m-%d %H:%M")}</div>' if metadata else ''}
</body>
</html>"""
            
            # Write to file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_doc)
            
            print(f"🌐 Exported HTML: {filename}")
            return str(file_path)
            
        except Exception as e:
            print(f"❌ HTML export error: {e}")
            raise
    
    async def _export_to_txt(self, content: str, filename: str, metadata: Optional[Dict] = None) -> str:
        """Export notes to plain text format"""
        
        try:
            file_path = self.export_dir / filename
            
            # Convert markdown to plain text
            plain_text = html2text.html2text(markdown.markdown(content))
            
            # Add metadata
            if metadata:
                plain_text += f"\n\nGenerated by StudyMate AI on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            
            # Write to file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(plain_text)
            
            print(f"📄 Exported TXT: {filename}")
            return str(file_path)
            
        except Exception as e:
            print(f"❌ TXT export error: {e}")
            raise
    
    async def _export_to_markdown(self, content: str, filename: str, metadata: Optional[Dict] = None) -> str:
        """Export notes to markdown format"""
        
        try:
            file_path = self.export_dir / filename
            
            # Add metadata
            final_content = content
            if metadata:
                final_content += f"\n\n---\n\n*Generated by StudyMate AI on {datetime.now().strftime('%Y-%m-%d %H:%M')}*"
            
            # Write to file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(final_content)
            
            print(f"📝 Exported Markdown: {filename}")
            return str(file_path)
            
        except Exception as e:
            print(f"❌ Markdown export error: {e}")
            raise
    
    async def generate_quiz_from_notes(self, notes_content: str, gemini_model = None) -> str:
        """Generate quiz questions from notes content"""
        
        if not gemini_model:
            return "Quiz generation requires AI capabilities. Please configure your Gemini API key."
        
        try:
            prompt = f"""Based on the following study notes, generate a comprehensive quiz with 10 questions:

{notes_content[:3000]}  # Limit content length

Create:
1. 5 multiple choice questions (4 options each)
2. 3 short answer questions
3. 2 essay questions

Format each question clearly with:
- Question number and type
- The question text
- For multiple choice: A, B, C, D options with correct answer marked
- For short answer: Expected key points
- For essay: Evaluation criteria

Make questions that test understanding, not just memorization."""

            response = gemini_model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            return f"Error generating quiz: {str(e)}"
    
    async def create_flashcards(self, notes_content: str, gemini_model = None) -> List[Dict[str, str]]:
        """Generate flashcards from notes content"""
        
        if not gemini_model:
            return [{"front": "Flashcard generation requires AI", "back": "Please configure your Gemini API key"}]
        
        try:
            prompt = f"""Create 15 flashcards from these study notes:

{notes_content[:2000]}

Format as JSON array with objects containing 'front' and 'back' fields.
Make flashcards that help with memorization and understanding.
Include key terms, definitions, concepts, and important facts.

Example format:
[
  {{"front": "What is...?", "back": "The answer is..."}},
  {{"front": "Define...", "back": "Definition..."}}
]"""

            response = gemini_model.generate_content(prompt)
            
            # Try to parse JSON response
            try:
                import json
                flashcards = json.loads(response.text)
                return flashcards if isinstance(flashcards, list) else []
            except:
                # Fallback: parse manually
                return self._parse_flashcards_from_text(response.text)
                
        except Exception as e:
            return [{"front": "Error generating flashcards", "back": str(e)}]
    
    def _parse_flashcards_from_text(self, text: str) -> List[Dict[str, str]]:
        """Parse flashcards from text when JSON parsing fails"""
        
        flashcards = []
        lines = text.split('\n')
        
        current_front = ""
        current_back = ""
        
        for line in lines:
            line = line.strip()
            if line.startswith('Q:') or line.startswith('Front:'):
                current_front = line.split(':', 1)[1].strip()
            elif line.startswith('A:') or line.startswith('Back:'):
                current_back = line.split(':', 1)[1].strip()
                if current_front and current_back:
                    flashcards.append({"front": current_front, "back": current_back})
                    current_front = ""
                    current_back = ""
        
        return flashcards[:15]  # Limit to 15 flashcards